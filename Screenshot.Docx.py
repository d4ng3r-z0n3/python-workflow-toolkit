#!/usr/bin/env python3

# For Debian/Ubuntu:
# sudo apt update
# sudo apt install scrot xclip

# Alternative screenshot tools you could use:
# sudo apt install gnome-screenshot
# sudo apt install spectacle
# sudo apt install flameshot

import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
import pyautogui
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import io
import os
import time
import json
from datetime import datetime
import requests
import hashlib
import subprocess
import sys
import threading
import socket
import webbrowser
import urllib.parse
import platform
import tempfile

CURRENT_OS = platform.system().lower()
is_windows = CURRENT_OS == "windows"
is_linux = CURRENT_OS == "linux" 
is_macos = CURRENT_OS == "darwin"

if is_windows:
    try:
        import ctypes
        import ctypes.wintypes as wt
        user32 = ctypes.windll.user32
        gdi32 = ctypes.windll.gdi32
        dwmapi = ctypes.windll.dwmapi
        try:
            user32.SetProcessDPIAware()
        except Exception:
            pass
        SRCCOPY = 0x00CC0020
        CAPTUREBLT = 0x40000000
        PW_RENDERFULLCONTENT = 0x00000002
        DWMWA_EXTENDED_FRAME_BOUNDS = 9
    except ImportError:
        print("Warning: Windows ctypes modules not available")
        is_windows = False

if is_linux:
    try:
        subprocess.run(['which', 'scrot'], check=True, capture_output=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("Warning: scrot not found - install with: sudo apt install scrot")

class LicenseDialog:
    def __init__(self, parent):
        self.result = None
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Student PenTester LLC - License Registration")
        self.dialog.geometry("600x650")
        self.dialog.resizable(True, False)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.dialog.protocol("WM_DELETE_WINDOW", self.on_cancel)
        
        self.create_license_dialog()
        
    def create_license_dialog(self):
        main_frame = ttk.Frame(self.dialog, padding=20)
        main_frame.pack(fill='both', expand=True)
        
        title_label = ttk.Label(main_frame, text="Student PenTester LLC", font=('Segoe UI', 16, 'bold'))
        title_label.pack(pady=(0, 10))
        
        subtitle_label = ttk.Label(main_frame, text="Screenshot to DOCX Generator", font=('Segoe UI', 12))
        subtitle_label.pack(pady=(0, 20))
        
        license_frame = ttk.LabelFrame(main_frame, text="GNU General Public License v3.0", padding=10)
        license_frame.pack(fill='both', expand=True, pady=(0, 20))
        
        license_text = tk.Text(license_frame, height=12, wrap='word', font=('Courier', 9))
        scrollbar = ttk.Scrollbar(license_frame, orient='vertical', command=license_text.yview)
        license_text.configure(yscrollcommand=scrollbar.set)
        
        license_content = """GNU GENERAL PUBLIC LICENSE Version 3, 29 June 2007

Copyright (C) 2025 Student PenTester LLC

This program is free software: you can redistribute it and/or modify it under the terms of the GNU General Public License as published by the Free Software Foundation, either version 3 of the License, or (at your option) any later version.

This program is distributed in the hope that it will be useful, but WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.

You should have received a copy of the GNU General Public License along with this program. If not, see <https://www.gnu.org/licenses/>.

By using this software, you acknowledge that Student PenTester LLC may collect usage statistics, update information, and contact details for support and improvement purposes."""
        
        license_text.insert('1.0', license_content)
        license_text.config(state='disabled')
        
        license_text.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        info_frame = ttk.LabelFrame(main_frame, text="Registration Information", padding=15)
        info_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(info_frame, text="Full Name:*").grid(row=0, column=0, sticky='w', pady=5)
        self.name_entry = ttk.Entry(info_frame, font=('Segoe UI', 10), width=30)
        self.name_entry.grid(row=0, column=1, sticky='ew', padx=(10, 0), pady=5)
        
        ttk.Label(info_frame, text="Email Address:*").grid(row=1, column=0, sticky='w', pady=5)
        self.email_entry = ttk.Entry(info_frame, font=('Segoe UI', 10), width=30)
        self.email_entry.grid(row=1, column=1, sticky='ew', padx=(10, 0), pady=5)
        
        info_frame.columnconfigure(1, weight=1)
        
        self.agree_var = tk.BooleanVar()
        agree_check = ttk.Checkbutton(main_frame, text="I agree to the terms of the GNU General Public License v3.0", variable=self.agree_var)
        agree_check.pack(pady=(0, 10))
        
        self.updates_var = tk.BooleanVar(value=True)
        updates_check = ttk.Checkbutton(main_frame, text="Enable automatic updates (recommended)", variable=self.updates_var)
        updates_check.pack(pady=(0, 20))
        
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill='x')
        
        accept_btn = ttk.Button(button_frame, text="Accept & Continue", command=self.on_accept)
        accept_btn.pack(side='right', padx=(10, 0), pady=5)
        
        cancel_btn = ttk.Button(button_frame, text="Cancel", command=self.on_cancel)
        cancel_btn.pack(side='right', pady=5)
        
        self.name_entry.focus()
        
    def on_accept(self):
        name = self.name_entry.get().strip()
        email = self.email_entry.get().strip()
        
        if not name:
            messagebox.showerror("Error", "Full name is required.")
            return
            
        if not email or '@' not in email:
            messagebox.showerror("Error", "Valid email address is required.")
            return
            
        if not self.agree_var.get():
            messagebox.showerror("Error", "You must agree to the license terms to continue.")
            return
            
        self.result = {
            'name': name,
            'email': email,
            'auto_updates': self.updates_var.get(),
            'accepted': True
        }
        self.dialog.destroy()
        
    def on_cancel(self):
        self.result = {'accepted': False}
        self.dialog.destroy()

class UpdateChecker:
    def __init__(self, app_instance):
        self.app = app_instance
        self.current_version = "1.0.3"
        self.update_server = "https://update.xn--mdaa.com"
        self.update_endpoint = f"{self.update_server}/api/check-update"
        self.download_endpoint = f"{self.update_server}/api/download"
        self.register_endpoint = f"{self.update_server}/api/register"

    def get_system_info(self):
        try:
            hostname = socket.gethostname()
            ip_address = socket.gethostbyname(hostname)
        except:
            hostname = "unknown"
            ip_address = "unknown"
            
        return {
            'hostname': hostname,
            'os': platform.system(),
            'python_version': sys.version,
            'platform': platform.platform()
        }
    
    def register_user(self, user_info):
        try:
            system_info = self.get_system_info()
            registration_data = {
                'name': user_info['name'],
                'email': user_info['email'],
                'version': self.current_version,
                'auto_updates': user_info['auto_updates'],
                'system_info': system_info,
                'timestamp': datetime.now().isoformat()
            }
            
            response = requests.post(self.register_endpoint, json=registration_data, timeout=10)
            return response.status_code == 200
        except:
            return True

    def check_for_updates(self, silent=False):
        if not self.app.settings.get('auto_updates', True) and silent:
            return
            
        def update_check_thread():
            try:
                system_info = self.get_system_info()
                check_data = {
                    'current_version': self.current_version,
                    'email': self.app.settings.get('email', ''),
                    'system_info': system_info,
                    'timestamp': datetime.now().isoformat()
                }
                
                response = requests.post(self.update_endpoint, json=check_data, timeout=15)
                
                if response.status_code == 200:
                    update_info = response.json()
                    if update_info.get('update_available', False):
                        self.app.root.after(0, lambda: self.prompt_update(update_info))
                    else:
                        if not silent:
                            self.app.root.after(0, lambda: messagebox.showinfo("Updates", "You have the latest version!"))
                else:
                    if not silent:
                        self.app.root.after(0, lambda: messagebox.showerror("Update Error", f"Update check failed with code {response.status_code}"))
                        
            except Exception as e:
                if not silent:
                    self.app.root.after(0, lambda: messagebox.showerror("Update Error", f"Update check failed: {str(e)}"))
        
        thread = threading.Thread(target=update_check_thread, daemon=True)
        thread.start()

    def prompt_update(self, update_info):
        new_version = update_info.get('version', 'Unknown')
        changelog = update_info.get('changelog', 'No changelog available.')
        
        dialog = tk.Toplevel(self.app.root)
        dialog.title("Update Available")
        dialog.geometry("500x400")
        dialog.transient(self.app.root)
        dialog.grab_set()
        
        main_frame = ttk.Frame(dialog, padding=20)
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text=f"Update Available: Version {new_version}", font=('Segoe UI', 14, 'bold')).pack(pady=(0, 10))
        ttk.Label(main_frame, text=f"Current Version: {self.current_version}").pack(pady=(0, 20))
        
        changelog_frame = ttk.LabelFrame(main_frame, text="What's New", padding=10)
        changelog_frame.pack(fill='both', expand=True, pady=(0, 20))
        
        changelog_text = tk.Text(changelog_frame, height=10, wrap='word')
        changelog_scroll = ttk.Scrollbar(changelog_frame, orient='vertical', command=changelog_text.yview)
        changelog_text.configure(yscrollcommand=changelog_scroll.set)
        
        changelog_text.insert('1.0', changelog)
        changelog_text.config(state='disabled')
        
        changelog_text.pack(side='left', fill='both', expand=True)
        changelog_scroll.pack(side='right', fill='y')
        
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill='x')
        
        ttk.Button(button_frame, text="Later", command=dialog.destroy).pack(side='right')

class DocxScreenshotApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Screenshot to DOCX Generator")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 700)
        
        self.screenshots = []
        self.section_names = []
        self.notes = []
        self.current_index = 0
        
        self.setup_styles()
        
        if not self.check_license():
            return
            
        self.load_settings()
        self.create_menu()
        self.create_widgets()
        
        self.updater = UpdateChecker(self)
        
        if self.settings.get('auto_updates', True):
            self.root.after(2000, lambda: self.updater.check_for_updates(silent=True))
        
    def check_license(self):
        license_file = 'license.json'
        
        if os.path.exists(license_file):
            try:
                with open(license_file, 'r') as f:
                    license_data = json.load(f)
                
                if license_data.get('accepted', False):
                    return True
                    
            except (json.JSONDecodeError, Exception):
                pass
        
        license_dialog = LicenseDialog(self.root)
        self.root.wait_window(license_dialog.dialog)
        
        if not license_dialog.result or not license_dialog.result.get('accepted', False):
            self.root.quit()
            return False
        
        try:
            with open(license_file, 'w') as f:
                json.dump(license_dialog.result, f, indent=2)
        except Exception:
            pass
        
        updater = UpdateChecker(self)
        updater.register_user(license_dialog.result)
        
        return True
        
    def setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        style.configure('Title.TLabel', font=('Segoe UI', 16, 'bold'), foreground='#2c3e50')
        style.configure('Subtitle.TLabel', font=('Segoe UI', 10), foreground='#34495e')
        style.configure('Action.TButton', font=('Segoe UI', 10, 'bold'), padding=(20, 10))
        style.configure('Small.TButton', font=('Segoe UI', 9), padding=(10, 5))
        
    def create_menu(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="About", command=self.show_about)
        help_menu.add_command(label="User Guide", command=self.show_help)
        help_menu.add_command(label="Check for Updates", command=lambda: self.updater.check_for_updates(silent=False))
        help_menu.add_separator()
        help_menu.add_command(label="Report Bug", command=self.report_bug)
        
        license_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="License", menu=license_menu)
        license_menu.add_command(label="View License", command=self.show_license)
        
        self.auto_updates_var = tk.BooleanVar(value=self.settings.get('auto_updates', True))
        license_menu.add_checkbutton(label="Auto Updates", variable=self.auto_updates_var, command=self.toggle_auto_updates)
        
        license_menu.add_separator()
        license_menu.add_command(label="Registration Info", command=self.show_registration)
        
    def show_about(self):
        about_text = """Professional Screenshot to DOCX Generator
Version 1.0.3

© 2025 Student PenTester LLC
Licensed under GNU General Public License v3.0

A powerful tool for capturing screenshots and generating professional DOCX documents with customizable formatting and organization features."""

        messagebox.showinfo("About", about_text)
    
    def show_help(self):
        help_window = tk.Toplevel(self.root)
        help_window.title("User Guide")
        help_window.geometry("600x500")
        help_window.transient(self.root)
        
        text_frame = ttk.Frame(help_window, padding=20)
        text_frame.pack(fill='both', expand=True)
        
        help_text = tk.Text(text_frame, wrap='word', font=('Segoe UI', 10))
        help_scroll = ttk.Scrollbar(text_frame, orient='vertical', command=help_text.yview)
        help_text.configure(yscrollcommand=help_scroll.set)
        
        help_content = """USER GUIDE

CAPTURE SCREENSHOTS:
1. Enter module number and document title
2. Set capture delay (1-10 seconds)
3. Click 'Capture Screenshot' and follow platform-specific instructions
4. Enter section name and optional notes for each screenshot
5. Use 'Import Image' to add existing images

EDIT & PREVIEW:
1. Select screenshots from the list to preview
2. Use arrow buttons to reorder screenshots
3. Edit section names and notes directly
4. Delete unwanted screenshots with trash button

NOTES FEATURE:
- Add notes below each screenshot for additional context
- Notes appear in the DOCX document below the image
- Use notes for explanations, instructions, or observations

SETTINGS:
1. Configure personal information for document headers
2. Set default save location
3. Adjust page margins and image heights
4. Save settings for future use

GENERATE DOCX:
1. Review all screenshots in Edit tab
2. Click 'Generate DOCX' in Capture tab
3. Document will be saved with timestamp
4. Choose to open file automatically

CROSS-PLATFORM SUPPORT:
- Windows: Advanced window capture
- Linux: Interactive selection with scrot
- macOS: Built-in screencapture utility
- All platforms: Fallback to general screenshot"""

        help_text.insert('1.0', help_content)
        help_text.config(state='disabled')
        
        help_text.pack(side='left', fill='both', expand=True)
        help_scroll.pack(side='right', fill='y')
    
    def report_bug(self):
        bug_window = tk.Toplevel(self.root)
        bug_window.title("Report Bug")
        bug_window.geometry("500x400")
        bug_window.transient(self.root)
        bug_window.grab_set()
        
        frame = ttk.Frame(bug_window, padding=20)
        frame.pack(fill='both', expand=True)
        
        ttk.Label(frame, text="Report a Bug", font=('Segoe UI', 14, 'bold')).pack(pady=(0, 10))
        ttk.Label(frame, text="Please describe the issue you encountered:").pack(anchor='w', pady=(0, 5))
        
        text_area = tk.Text(frame, height=10, wrap='word')
        text_area.pack(fill='both', expand=True, pady=(0, 10))
        
        button_frame = ttk.Frame(frame)
        button_frame.pack(fill='x')
        
        def submit_bug():
            bug_report = text_area.get('1.0', 'end-1c').strip()
            if bug_report:
                try:
                    system_info = self.updater.get_system_info()
                    report_data = {
                        'bug_report': bug_report,
                        'version': self.updater.current_version,
                        'email': self.settings.get('email', ''),
                        'system_info': system_info,
                        'timestamp': datetime.now().isoformat()
                    }
                    
                    response = requests.post(f"{self.updater.update_server}/api/bug-report", json=report_data, timeout=10)
                    if response.status_code == 200:
                        messagebox.showinfo("Bug Report", "Bug report submitted successfully. Thank you!")
                        bug_window.destroy()
                    else:
                        messagebox.showerror("Error", "Failed to submit bug report. Please try again later.")
                except:
                    messagebox.showerror("Error", "Unable to connect to server. Please try again later.")
            else:
                messagebox.showwarning("Warning", "Please describe the bug before submitting.")
        
        ttk.Button(button_frame, text="Submit", command=submit_bug).pack(side='right', padx=(10, 0))
        ttk.Button(button_frame, text="Cancel", command=bug_window.destroy).pack(side='right')
    
    def show_license(self):
        license_window = tk.Toplevel(self.root)
        license_window.title("GNU General Public License v3.0")
        license_window.geometry("600x500")
        license_window.transient(self.root)
        
        text_frame = ttk.Frame(license_window, padding=20)
        text_frame.pack(fill='both', expand=True)
        
        license_text = tk.Text(text_frame, wrap='word', font=('Courier', 9))
        license_scroll = ttk.Scrollbar(text_frame, orient='vertical', command=license_text.yview)
        license_text.configure(yscrollcommand=license_scroll.set)
        
        license_content = """GNU GENERAL PUBLIC LICENSE
Version 3, 29 June 2007

Copyright (C) 2025 Student PenTester LLC

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <https://www.gnu.org/licenses/>."""
        
        license_text.insert('1.0', license_content)
        license_text.config(state='disabled')
        
        license_text.pack(side='left', fill='both', expand=True)
        license_scroll.pack(side='right', fill='y')
    
    def toggle_auto_updates(self):
        new_setting = self.auto_updates_var.get()
        self.settings['auto_updates'] = new_setting
        self.save_settings()
        
        status = "enabled" if new_setting else "disabled"
        messagebox.showinfo("Auto Updates", f"Automatic updates have been {status}.")
    
    def show_registration(self):
        try:
            with open('license.json', 'r') as f:
                license_data = json.load(f)
        except:
            license_data = {}
        
        reg_info = f"""Registration Information:

Name: {license_data.get('name', 'Unknown')}
Email: {license_data.get('email', 'Unknown')}
Auto Updates: {'Enabled' if self.settings.get('auto_updates', True) else 'Disabled'}
License: GNU GPL v3.0
Registered: Student PenTester LLC"""

        messagebox.showinfo("Registration Info", reg_info)
        
    def load_settings(self):
        try:
            with open('screenshot_app_settings.json', 'r') as f:
                self.settings = json.load(f)
        except:
            self.settings = {}
            
        try:
            with open('license.json', 'r') as f:
                license_data = json.load(f)
                self.settings['email'] = license_data.get('email', '')
                if 'auto_updates' not in self.settings:
                    self.settings['auto_updates'] = license_data.get('auto_updates', True)
        except:
            pass
            
        self.first_name = self.settings.get('first_name', 'First Name')
        self.last_name = self.settings.get('last_name', 'Last Name')
        self.course_code = self.settings.get('course_code', 'COURSE001')
        self.default_save_path = self.settings.get('save_path', os.path.join(os.path.expanduser("~"), "Documents"))
    
    def save_settings(self):
        try:
            with open('screenshot_app_settings.json', 'w') as f:
                json.dump(self.settings, f)
        except:
            pass
    
    def create_widgets(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        self.capture_frame = ttk.Frame(notebook)
        self.edit_frame = ttk.Frame(notebook)
        self.settings_frame = ttk.Frame(notebook)
        
        notebook.add(self.capture_frame, text="Capture Screenshots")
        notebook.add(self.edit_frame, text="Edit & Preview")
        notebook.add(self.settings_frame, text="Settings")
        
        self.create_capture_tab()
        self.create_edit_tab()
        self.create_settings_tab()
        
    def create_capture_tab(self):
        header_frame = ttk.Frame(self.capture_frame)
        header_frame.pack(fill='x', padx=20, pady=20)
        
        ttk.Label(header_frame, text="Screenshot Capture", style='Title.TLabel').pack()
        ttk.Label(header_frame, text="Capture screenshots and organize them for document generation", style='Subtitle.TLabel').pack(pady=(5, 0))
        
        input_frame = ttk.LabelFrame(self.capture_frame, text="Document Information", padding=20)
        input_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Label(input_frame, text="Module/Assignment #:").grid(row=0, column=0, sticky='w', padx=(0, 10))
        self.module_entry = ttk.Entry(input_frame, font=('Segoe UI', 10))
        self.module_entry.grid(row=0, column=1, sticky='ew', padx=(0, 20))
        
        ttk.Label(input_frame, text="Document Title:").grid(row=0, column=2, sticky='w', padx=(0, 10))
        self.doc_title_entry = ttk.Entry(input_frame, font=('Segoe UI', 10))
        self.doc_title_entry.grid(row=0, column=3, sticky='ew')
        self.doc_title_entry.insert(0, "Interactive Sections")
        
        input_frame.columnconfigure(1, weight=1)
        input_frame.columnconfigure(3, weight=1)
        
        capture_frame = ttk.LabelFrame(self.capture_frame, text="Capture Options", padding=20)
        capture_frame.pack(fill='x', padx=20, pady=10)
        
        self.capture_delay = tk.IntVar(value=3)
        ttk.Label(capture_frame, text="Capture Delay (seconds):").pack(side='left')
        delay_spin = ttk.Spinbox(capture_frame, from_=1, to=10, textvariable=self.capture_delay, width=5)
        delay_spin.pack(side='left', padx=(10, 20))
        
        ttk.Button(capture_frame, text="Capture Screenshot", style='Action.TButton', command=self.capture_screenshot).pack(side='left', padx=10)
        ttk.Button(capture_frame, text="Import Image", style='Small.TButton', command=self.import_image).pack(side='left')
        
        section_input_frame = ttk.LabelFrame(self.capture_frame, text="Section Name for Next Screenshot", padding=15)
        section_input_frame.pack(fill='x', padx=20, pady=10)
        
        section_label_frame = ttk.Frame(section_input_frame)
        section_label_frame.pack(fill='x', pady=(0, 5))
        
        ttk.Label(section_label_frame, text="Section Name:").pack(side='left', padx=(0, 10))
        
        info_label = ttk.Label(section_label_frame, text="If no entry here, you will be asked to name it in popup window", 
                              font=('Segoe UI', 8, 'italic'), foreground='red')
        info_label.pack(side='left')
        
        self.section_entry = ttk.Entry(section_input_frame, font=('Segoe UI', 10))
        self.section_entry.pack(fill='x', expand=True, pady=(0, 10))
        
        ttk.Label(section_input_frame, text="Notes (optional):").pack(anchor='w')
        self.notes_entry = tk.Text(section_input_frame, height=3, wrap='word', font=('Segoe UI', 9))
        self.notes_entry.pack(fill='x', expand=True)
        
        status_frame = ttk.Frame(self.capture_frame)
        status_frame.pack(fill='x', padx=20, pady=10)
        
        self.status_label = ttk.Label(status_frame, text="No screenshots captured", font=('Segoe UI', 10))
        self.status_label.pack()
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(status_frame, variable=self.progress_var, mode='determinate')
        
        action_frame = ttk.Frame(self.capture_frame)
        action_frame.pack(fill='x', padx=20, pady=20)
        
        ttk.Button(action_frame, text="Generate DOCX", style='Action.TButton', command=self.generate_docx).pack(side='left', padx=(0, 10))
        ttk.Button(action_frame, text="Save Project", style='Small.TButton', command=self.save_project).pack(side='left', padx=(0, 10))
        ttk.Button(action_frame, text="Load Project", style='Small.TButton', command=self.load_project).pack(side='left')
        
    def create_edit_tab(self):
        header_frame = ttk.Frame(self.edit_frame)
        header_frame.pack(fill='x', padx=20, pady=20)
        
        ttk.Label(header_frame, text="Edit & Preview", style='Title.TLabel').pack()
        ttk.Label(header_frame, text="Review, edit, and organize your screenshots before generating the document", style='Subtitle.TLabel').pack(pady=(5, 0))
        
        content_frame = ttk.PanedWindow(self.edit_frame, orient='horizontal')
        content_frame.pack(fill='both', expand=True, padx=20, pady=10)
        
        left_panel = ttk.Frame(content_frame)
        right_panel = ttk.Frame(content_frame)
        content_frame.add(left_panel, weight=1)
        content_frame.add(right_panel, weight=2)
        
        list_frame = ttk.LabelFrame(left_panel, text="Screenshots List", padding=10)
        list_frame.pack(fill='both', expand=True)
        
        list_scroll_frame = ttk.Frame(list_frame)
        list_scroll_frame.pack(fill='both', expand=True)
        
        list_scrollbar = ttk.Scrollbar(list_scroll_frame)
        list_scrollbar.pack(side='right', fill='y')
        
        self.screenshots_listbox = tk.Listbox(list_scroll_frame, yscrollcommand=list_scrollbar.set, font=('Segoe UI', 10))
        self.screenshots_listbox.pack(side='left', fill='both', expand=True)
        self.screenshots_listbox.bind('<<ListboxSelect>>', self.on_screenshot_select)
        list_scrollbar.config(command=self.screenshots_listbox.yview)
        
        list_buttons_frame = ttk.Frame(list_frame)
        list_buttons_frame.pack(fill='x', pady=(10, 0))
        
        ttk.Button(list_buttons_frame, text="↑", command=self.move_up, width=3).pack(side='left', padx=(0, 5))
        ttk.Button(list_buttons_frame, text="↓", command=self.move_down, width=3).pack(side='left', padx=(0, 5))
        ttk.Button(list_buttons_frame, text="Del", command=self.delete_screenshot, width=3).pack(side='left', padx=(0, 5))
        ttk.Button(list_buttons_frame, text="Edit", command=self.edit_section_name, width=3).pack(side='left')
        
        preview_frame = ttk.LabelFrame(right_panel, text="Preview", padding=10)
        preview_frame.pack(fill='both', expand=True)
        
        self.canvas_frame = ttk.Frame(preview_frame)
        self.canvas_frame.pack(fill='both', expand=True)
        
        self.canvas = tk.Canvas(self.canvas_frame, bg='white')
        canvas_h_scroll = ttk.Scrollbar(self.canvas_frame, orient='horizontal', command=self.canvas.xview)
        canvas_v_scroll = ttk.Scrollbar(self.canvas_frame, orient='vertical', command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=canvas_h_scroll.set, yscrollcommand=canvas_v_scroll.set)
        
        canvas_h_scroll.pack(side='bottom', fill='x')
        canvas_v_scroll.pack(side='right', fill='y')
        self.canvas.pack(side='left', fill='both', expand=True)
        
        preview_controls = ttk.Frame(preview_frame)
        preview_controls.pack(fill='x', pady=(10, 0))
        
        ttk.Label(preview_controls, text="Section Name:").pack(side='left')
        self.preview_section_entry = ttk.Entry(preview_controls, font=('Segoe UI', 10))
        self.preview_section_entry.pack(side='left', fill='x', expand=True, padx=(10, 10))
        ttk.Button(preview_controls, text="Update", command=self.update_section_name).pack(side='left')
        
        notes_control_frame = ttk.Frame(preview_frame)
        notes_control_frame.pack(fill='both', expand=True, pady=(10, 0))
        
        ttk.Label(notes_control_frame, text="Notes:").pack(anchor='w')
        
        notes_text_frame = ttk.Frame(notes_control_frame)
        notes_text_frame.pack(fill='both', expand=True)
        
        self.preview_notes_text = tk.Text(notes_text_frame, height=4, wrap='word', font=('Segoe UI', 9))
        notes_scroll = ttk.Scrollbar(notes_text_frame, orient='vertical', command=self.preview_notes_text.yview)
        self.preview_notes_text.configure(yscrollcommand=notes_scroll.set)
        
        self.preview_notes_text.pack(side='left', fill='both', expand=True)
        notes_scroll.pack(side='right', fill='y')
        
        notes_button_frame = ttk.Frame(notes_control_frame)
        notes_button_frame.pack(fill='x', pady=(5, 0))
        
        ttk.Button(notes_button_frame, text="Update Notes", command=self.update_notes).pack(side='right')
        
    def create_settings_tab(self):
        header_frame = ttk.Frame(self.settings_frame)
        header_frame.pack(fill='x', padx=20, pady=20)
        
        ttk.Label(header_frame, text="Settings", style='Title.TLabel').pack()
        ttk.Label(header_frame, text="Configure your personal information and preferences", style='Subtitle.TLabel').pack(pady=(5, 0))
        
        personal_frame = ttk.LabelFrame(self.settings_frame, text="Personal Information", padding=20)
        personal_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Label(personal_frame, text="First Name:").grid(row=0, column=0, sticky='w', padx=(0, 10), pady=5)
        self.first_name_entry = ttk.Entry(personal_frame, font=('Segoe UI', 10))
        self.first_name_entry.grid(row=0, column=1, sticky='ew', pady=5)
        self.first_name_entry.insert(0, self.first_name)
        
        ttk.Label(personal_frame, text="Last Name:").grid(row=1, column=0, sticky='w', padx=(0, 10), pady=5)
        self.last_name_entry = ttk.Entry(personal_frame, font=('Segoe UI', 10))
        self.last_name_entry.grid(row=1, column=1, sticky='ew', pady=5)
        self.last_name_entry.insert(0, self.last_name)
        
        ttk.Label(personal_frame, text="Course Code:").grid(row=2, column=0, sticky='w', padx=(0, 10), pady=5)
        self.course_code_entry = ttk.Entry(personal_frame, font=('Segoe UI', 10))
        self.course_code_entry.grid(row=2, column=1, sticky='ew', pady=5)
        self.course_code_entry.insert(0, self.course_code)
        
        personal_frame.columnconfigure(1, weight=1)
        
        save_frame = ttk.LabelFrame(self.settings_frame, text="Save Location", padding=20)
        save_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Label(save_frame, text="Default Save Path:").pack(anchor='w')
        path_frame = ttk.Frame(save_frame)
        path_frame.pack(fill='x', pady=(5, 0))
        
        self.save_path_entry = ttk.Entry(path_frame, font=('Segoe UI', 10))
        self.save_path_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))
        self.save_path_entry.insert(0, self.default_save_path)
        ttk.Button(path_frame, text="Browse", command=self.browse_save_path).pack(side='right')
        
        format_frame = ttk.LabelFrame(self.settings_frame, text="Document Format", padding=20)
        format_frame.pack(fill='x', padx=20, pady=10)
        
        self.margin_var = tk.DoubleVar(value=0.25)
        ttk.Label(format_frame, text="Page Margins (inches):").pack(anchor='w')
        margin_frame = ttk.Frame(format_frame)
        margin_frame.pack(fill='x', pady=(5, 0))
        ttk.Spinbox(margin_frame, from_=0.1, to=2.0, increment=0.25, textvariable=self.margin_var, width=10).pack(side='left')
        
        self.image_height_var = tk.DoubleVar(value=6.5)
        ttk.Label(format_frame, text="Image Height (inches):").pack(anchor='w', pady=(10, 0))
        height_frame = ttk.Frame(format_frame)
        height_frame.pack(fill='x', pady=(5, 0))
        ttk.Spinbox(height_frame, from_=3.0, to=10.0, increment=0.5, textvariable=self.image_height_var, width=10).pack(side='left')
        
        buttons_frame = ttk.Frame(self.settings_frame)
        buttons_frame.pack(fill='x', padx=20, pady=20)
        
        ttk.Button(buttons_frame, text="Save Settings", style='Action.TButton', command=self.apply_settings).pack(side='left', padx=(0, 10))
        ttk.Button(buttons_frame, text="Reset to Default", style='Small.TButton', command=self.reset_settings).pack(side='left')

    def _get_window_rect(self, hwnd):
        if not is_windows:
            return None
        rect = wt.RECT()
        ok = dwmapi.DwmGetWindowAttribute(wt.HWND(hwnd), wt.DWORD(DWMWA_EXTENDED_FRAME_BOUNDS), ctypes.byref(rect), ctypes.sizeof(rect))
        if ok != 0:
            user32.GetWindowRect(wt.HWND(hwnd), ctypes.byref(rect))
        return rect.left, rect.top, rect.right, rect.bottom

    def _pil_image_from_hbitmap(self, hdc, hbmp, width, height):
        if not is_windows:
            return None
        bmi = ctypes.create_string_buffer(40 + 4 * 256)
        ctypes.memset(bmi, 0, 40)
        ctypes.cast(bmi, ctypes.POINTER(ctypes.c_uint32))[0] = 40
        ctypes.cast(bmi, ctypes.POINTER(ctypes.c_int32))[1] = width
        ctypes.cast(bmi, ctypes.POINTER(ctypes.c_int32))[2] = -height
        ctypes.cast(bmi, ctypes.POINTER(ctypes.c_uint16))[6] = 1
        ctypes.cast(bmi, ctypes.POINTER(ctypes.c_uint16))[7] = 32
        ctypes.cast(bmi, ctypes.POINTER(ctypes.c_uint32))[5] = 0
        buf_len = width * height * 4
        pixel_data = ctypes.create_string_buffer(buf_len)
        gdi32.GetDIBits(hdc, hbmp, 0, height, pixel_data, bmi, 0)
        img = Image.frombuffer("RGBA", (width, height), pixel_data, "raw", "BGRA", 0, 1)
        return img

    def _capture_window_windows(self, timeout=3):
        if not is_windows:
            return None
        try:
            self.root.iconify()
            messagebox.showinfo("Capture", f"Hover mouse over target window; capturing in {timeout}s.")
            time.sleep(timeout)
            pt = wt.POINT()
            user32.GetCursorPos(ctypes.byref(pt))
            hwnd = user32.WindowFromPoint(pt)
            if not hwnd or user32.IsIconic(hwnd):
                return None
            l, t, r, b = self._get_window_rect(hwnd)
            width, height = r - l, b - t
            if width <= 0 or height <= 0:
                return None
            hdc_window = user32.GetWindowDC(hwnd)
            memdc = gdi32.CreateCompatibleDC(hdc_window)
            hbmp = gdi32.CreateCompatibleBitmap(hdc_window, width, height)
            gdi32.SelectObject(memdc, hbmp)
            ok = ctypes.windll.user32.PrintWindow(hwnd, memdc, PW_RENDERFULLCONTENT)
            if ok == 0:
                ok = gdi32.BitBlt(memdc, 0, 0, width, height, hdc_window, 0, 0, SRCCOPY | CAPTUREBLT)
            img = None
            if ok != 0:
                img = self._pil_image_from_hbitmap(memdc, hbmp, width, height)
            gdi32.DeleteObject(hbmp)
            gdi32.DeleteDC(memdc)
            user32.ReleaseDC(hwnd, hdc_window)
            if img is None:
                img = pyautogui.screenshot(region=(l, t, width, height))
            return img
        except Exception:
            return None
        finally:
            try:
                self.root.deiconify()
            except Exception:
                pass

    def _capture_window_linux(self, timeout=3):
        try:
            self.root.iconify()
            messagebox.showinfo("Capture", f"Click on target window to capture in {timeout}s.")
            time.sleep(timeout)
            
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                result = subprocess.run(['scrot', '-s', tmp_path], 
                                      capture_output=True, 
                                      timeout=30)
                
                if result.returncode == 0 and os.path.exists(tmp_path):
                    img = Image.open(tmp_path)
                    os.unlink(tmp_path)
                    return img
                else:
                    return pyautogui.screenshot()
            except subprocess.TimeoutExpired:
                return pyautogui.screenshot()
            except Exception:
                return pyautogui.screenshot()
                
        except Exception:
            return pyautogui.screenshot()
        finally:
            try:
                self.root.deiconify()
            except Exception:
                pass

    def _capture_window_macos(self, timeout=3):
        try:
            self.root.iconify()
            messagebox.showinfo("Capture", f"Screenshot will be taken in {timeout}s. Use Cmd+Shift+4 for selection.")
            time.sleep(timeout)
            
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                result = subprocess.run(['screencapture', '-i', tmp_path], 
                                      capture_output=True, 
                                      timeout=30)
                
                if result.returncode == 0 and os.path.exists(tmp_path):
                    img = Image.open(tmp_path)
                    os.unlink(tmp_path)
                    return img
                else:
                    return pyautogui.screenshot()
            except subprocess.TimeoutExpired:
                return pyautogui.screenshot()
            except Exception:
                return pyautogui.screenshot()
                
        except Exception:
            return pyautogui.screenshot()
        finally:
            try:
                self.root.deiconify()
            except Exception:
                pass

    def _capture_window_generic(self, timeout=3):
        try:
            self.root.iconify()
            messagebox.showinfo("Capture", f"Full screen capture in {timeout}s.")
            time.sleep(timeout)
            return pyautogui.screenshot()
        except Exception:
            return None
        finally:
            try:
                self.root.deiconify()
            except Exception:
                pass

    def capture_screenshot(self):
        timeout = self.capture_delay.get()
        img = None
        
        if is_windows:
            img = self._capture_window_windows(timeout)
        elif is_linux:
            img = self._capture_window_linux(timeout)
        elif is_macos:
            img = self._capture_window_macos(timeout)
        
        if img is None:
            img = self._capture_window_generic(timeout)
        
        if img is None:
            messagebox.showerror("Error", "Screenshot capture failed.")
            return
            
        try:
            section_name = self.section_entry.get().strip()
            if not section_name:
                section_name = simpledialog.askstring("Section Name", "Enter Section Name for this screenshot:", parent=self.root)
                if not section_name:
                    messagebox.showwarning("Warning", "Section name is required!")
                    return
            
            notes = self.notes_entry.get('1.0', 'end-1c').strip()
            
            self.screenshots.append(img)
            self.section_names.append(section_name)
            self.notes.append(notes)
            self.update_screenshot_list()
            self.status_label.config(text=f"Captured {len(self.screenshots)} screenshot(s)")
            
            self.section_entry.delete(0, 'end')
            self.notes_entry.delete('1.0', 'end')
            
        finally:
            try:
                self.root.deiconify()
            except Exception:
                pass

    def import_image(self):
        file_path = filedialog.askopenfilename(
            title="Select Image File",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"), ("All files", "*.*")]
        )
        if file_path:
            try:
                img = Image.open(file_path)
                section_name = self.section_entry.get().strip()
                if not section_name:
                    section_name = simpledialog.askstring("Section Name", "Enter Section Name for this image:", parent=self.root)
                    if not section_name:
                        messagebox.showwarning("Warning", "Section name is required!")
                        return
                
                notes = self.notes_entry.get('1.0', 'end-1c').strip()
                
                self.screenshots.append(img)
                self.section_names.append(section_name)
                self.notes.append(notes)
                self.update_screenshot_list()
                self.status_label.config(text=f"Captured {len(self.screenshots)} screenshot(s)")
                
                self.section_entry.delete(0, 'end')
                self.notes_entry.delete('1.0', 'end')
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load image: {str(e)}")

    def update_screenshot_list(self):
        self.screenshots_listbox.delete(0, 'end')
        for i, name in enumerate(self.section_names):
            self.screenshots_listbox.insert('end', f"{i+1}. {name}")

    def on_screenshot_select(self, event):
        selection = self.screenshots_listbox.curselection()
        if selection:
            self.current_index = selection[0]
            self.display_screenshot(self.current_index)

    def display_screenshot(self, index):
        if 0 <= index < len(self.screenshots):
            img = self.screenshots[index]
            
            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()
            
            if canvas_width > 1 and canvas_height > 1:
                img_ratio = img.width / img.height
                canvas_ratio = canvas_width / canvas_height
                
                if img_ratio > canvas_ratio:
                    display_width = min(canvas_width - 20, img.width)
                    display_height = int(display_width / img_ratio)
                else:
                    display_height = min(canvas_height - 20, img.height)
                    display_width = int(display_height * img_ratio)
                
                display_img = img.resize((display_width, display_height), Image.Resampling.LANCZOS)
                self.photo = ImageTk.PhotoImage(display_img)
                
                self.canvas.delete("all")
                self.canvas.create_image(canvas_width//2, canvas_height//2, image=self.photo)
                self.canvas.configure(scrollregion=self.canvas.bbox("all"))
            
            self.preview_section_entry.delete(0, 'end')
            self.preview_section_entry.insert(0, self.section_names[index])
            
            self.preview_notes_text.delete('1.0', 'end')
            if index < len(self.notes):
                self.preview_notes_text.insert('1.0', self.notes[index])

    def move_up(self):
        selection = self.screenshots_listbox.curselection()
        if selection and selection[0] > 0:
            index = selection[0]
            self.screenshots[index], self.screenshots[index-1] = self.screenshots[index-1], self.screenshots[index]
            self.section_names[index], self.section_names[index-1] = self.section_names[index-1], self.section_names[index]
            if len(self.notes) > index and len(self.notes) > index-1:
                self.notes[index], self.notes[index-1] = self.notes[index-1], self.notes[index]
            self.update_screenshot_list()
            self.screenshots_listbox.selection_set(index-1)

    def move_down(self):
        selection = self.screenshots_listbox.curselection()
        if selection and selection[0] < len(self.screenshots) - 1:
            index = selection[0]
            self.screenshots[index], self.screenshots[index+1] = self.screenshots[index+1], self.screenshots[index]
            self.section_names[index], self.section_names[index+1] = self.section_names[index+1], self.section_names[index]
            if len(self.notes) > index and len(self.notes) > index+1:
                self.notes[index], self.notes[index+1] = self.notes[index+1], self.notes[index]
            self.update_screenshot_list()
            self.screenshots_listbox.selection_set(index+1)

    def delete_screenshot(self):
        selection = self.screenshots_listbox.curselection()
        if selection:
            if messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this screenshot?"):
                index = selection[0]
                del self.screenshots[index]
                del self.section_names[index]
                if index < len(self.notes):
                    del self.notes[index]
                self.update_screenshot_list()
                self.canvas.delete("all")
                self.preview_section_entry.delete(0, 'end')
                self.preview_notes_text.delete('1.0', 'end')
                self.status_label.config(text=f"Captured {len(self.screenshots)} screenshot(s)")

    def edit_section_name(self):
        selection = self.screenshots_listbox.curselection()
        if selection:
            index = selection[0]
            new_name = simpledialog.askstring("Edit Section Name", "Enter new section name:", initialvalue=self.section_names[index])
            if new_name:
                self.section_names[index] = new_name
                self.update_screenshot_list()
                self.screenshots_listbox.selection_set(index)
                if hasattr(self, 'preview_section_entry'):
                    self.preview_section_entry.delete(0, 'end')
                    self.preview_section_entry.insert(0, new_name)

    def update_section_name(self):
        selection = self.screenshots_listbox.curselection()
        if selection:
            index = selection[0]
            new_name = self.preview_section_entry.get()
            if new_name:
                self.section_names[index] = new_name
                self.update_screenshot_list()
                self.screenshots_listbox.selection_set(index)

    def update_notes(self):
        selection = self.screenshots_listbox.curselection()
        if selection:
            index = selection[0]
            new_notes = self.preview_notes_text.get('1.0', 'end-1c')
            
            while len(self.notes) <= index:
                self.notes.append("")
            
            self.notes[index] = new_notes
            messagebox.showinfo("Notes Updated", "Notes have been updated successfully!")

    def apply_settings(self):
        self.first_name = self.first_name_entry.get()
        self.last_name = self.last_name_entry.get()
        self.course_code = self.course_code_entry.get()
        self.default_save_path = self.save_path_entry.get()
        
        self.settings.update({
            'first_name': self.first_name,
            'last_name': self.last_name,
            'course_code': self.course_code,
            'save_path': self.default_save_path
        })
        
        self.save_settings()
        messagebox.showinfo("Settings", "Settings saved successfully!")

    def browse_save_path(self):
        path = filedialog.askdirectory(initialdir=self.default_save_path)
        if path:
            self.save_path_entry.delete(0, 'end')
            self.save_path_entry.insert(0, path)

    def reset_settings(self):
        if messagebox.askyesno("Reset Settings", "Are you sure you want to reset all settings to default?"):
            self.first_name_entry.delete(0, 'end')
            self.first_name_entry.insert(0, "Your Name")
            self.last_name_entry.delete(0, 'end')
            self.last_name_entry.insert(0, "Last Name")
            self.course_code_entry.delete(0, 'end')
            self.course_code_entry.insert(0, "COURSE001")
            self.save_path_entry.delete(0, 'end')
            self.save_path_entry.insert(0, os.path.join(os.path.expanduser("~"), "Documents"))
            self.margin_var.set(0.25)
            self.image_height_var.set(6.5)

    def save_project(self):
        if not self.screenshots:
            messagebox.showwarning("Warning", "No screenshots to save!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".ssp",
            filetypes=[("Screenshot Project", "*.ssp"), ("All files", "*.*")],
            title="Save Screenshot Project"
        )
        
        if file_path:
            try:
                project_data = {
                    'section_names': self.section_names,
                    'notes': self.notes,
                    'module': self.module_entry.get(),
                    'doc_title': self.doc_title_entry.get(),
                    'created': datetime.now().isoformat(),
                    'screenshot_count': len(self.screenshots)
                }
                
                project_dir = file_path + "_data"
                os.makedirs(project_dir, exist_ok=True)
                
                for i, img in enumerate(self.screenshots):
                    img_path = os.path.join(project_dir, f"screenshot_{i}.png")
                    img.save(img_path, "PNG")
                
                with open(file_path, 'w') as f:
                    json.dump(project_data, f, indent=2)
                
                messagebox.showinfo("Success", f"Project saved successfully!")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save project: {str(e)}")

    def load_project(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Screenshot Project", "*.ssp"), ("All files", "*.*")],
            title="Load Screenshot Project"
        )
        
        if file_path:
            try:
                with open(file_path, 'r') as f:
                    project_data = json.load(f)
                
                project_dir = file_path + "_data"
                
                self.screenshots = []
                self.section_names = project_data.get('section_names', [])
                self.notes = project_data.get('notes', [])
                
                for i in range(project_data.get('screenshot_count', 0)):
                    img_path = os.path.join(project_dir, f"screenshot_{i}.png")
                    if os.path.exists(img_path):
                        img = Image.open(img_path)
                        self.screenshots.append(img)
                
                while len(self.notes) < len(self.screenshots):
                    self.notes.append("")
                
                self.module_entry.delete(0, 'end')
                self.module_entry.insert(0, project_data.get('module', ''))
                
                self.doc_title_entry.delete(0, 'end')
                self.doc_title_entry.insert(0, project_data.get('doc_title', 'Interactive Sections'))
                
                self.update_screenshot_list()
                self.status_label.config(text=f"Loaded {len(self.screenshots)} screenshot(s)")
                
                messagebox.showinfo("Success", "Project loaded successfully!")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load project: {str(e)}")

    def generate_docx(self):
        if not self.screenshots:
            messagebox.showerror("Error", "No screenshots captured!")
            return
        
        module = self.module_entry.get()
        if not module:
            messagebox.showerror("Error", "Module number is required!")
            return
        
        doc_title = self.doc_title_entry.get() or "Interactive Sections"
        
        try:
            self.progress_bar.pack(fill='x', pady=(10, 0))
            self.progress_var.set(0)
            self.root.update()
            
            doc = Document()
            section = doc.sections[0]
            
            margin = self.margin_var.get()
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{self.first_name} {self.last_name}   {self.course_code}   Module {module} {doc_title}"
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')
            if cols:
                cols[0].set('num', '1')
            
            total_screenshots = len(self.screenshots)
            image_height = self.image_height_var.get()
            
            for i, img in enumerate(self.screenshots):
                self.progress_var.set((i / total_screenshots) * 90)
                self.root.update()
                
                p = doc.add_paragraph(self.section_names[i])
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.space_before = Pt(6)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                img_stream = io.BytesIO()
                img.save(img_stream, format='PNG')
                img_stream.seek(0)
                
                pic = doc.add_picture(img_stream, height=Inches(image_height))
                img_stream.close()
                
                pic_paragraph = pic._inline.xpath('ancestor::w:p')[0]
                pic_paragraph.set(qn('w:jc'), 'center')
                
                if i < len(self.notes) and self.notes[i].strip():
                    notes_paragraph = doc.add_paragraph()
                    notes_run = notes_paragraph.add_run(self.notes[i])
                    notes_run.font.size = Pt(10)
                    notes_run.font.italic = True
                    notes_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    notes_paragraph.paragraph_format.space_after = Pt(12)
                    notes_paragraph.paragraph_format.space_before = Pt(6)
                    notes_paragraph.paragraph_format.left_indent = Inches(0.25)
                
                if i < len(self.screenshots) - 1:
                    doc.add_page_break()
            
            self.progress_var.set(95)
            self.root.update()
            
            # Fix: Save to current directory instead of Documents
            dir_path = os.getcwd()  # Current working directory
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{self.first_name.replace(' ', '.')}.{self.last_name.replace(' ', '.')}.Module{module}_{timestamp}.docx"
            full_path = os.path.join(dir_path, filename)
            
            doc.save(full_path)
            
            self.progress_var.set(100)
            self.root.update()
            
            messagebox.showinfo("Success", f"Document saved as {filename} in {dir_path}")
            
            if messagebox.askyesno("Open File", "Do you want to open the file?"):
                try:
                    if is_windows:
                        os.startfile(full_path)
                    elif is_macos:
                        subprocess.run(['open', full_path])
                    elif is_linux:
                        subprocess.run(['xdg-open', full_path])
                except Exception:
                    messagebox.showwarning("Open", "Unable to open file automatically.")
            
            if messagebox.askyesno("Clear Screenshots", "Do you want to clear all screenshots for a new project?"):
                self.screenshots = []
                self.section_names = []
                self.notes = []
                self.update_screenshot_list()
                self.canvas.delete("all")
                self.preview_section_entry.delete(0, 'end')
                self.preview_notes_text.delete('1.0', 'end')
                self.status_label.config(text="No screenshots captured")
                
        except PermissionError:
            messagebox.showerror("Error", f"Permission denied: Cannot save to {full_path}. Close the file if open and try again.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save document: {str(e)}")
        finally:
            self.progress_bar.pack_forget()
            self.progress_var.set(0)

if __name__ == "__main__":
    root = tk.Tk()
    app = DocxScreenshotApp(root)
    root.mainloop()
